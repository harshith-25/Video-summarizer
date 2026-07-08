import os
import logging
from datetime import datetime
from typing import Dict, Any
from app.config import Config

# python-docx imports
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ReportLab imports for premium PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

logger = logging.getLogger("app")

# Try to register Nirmala UI (Windows standard Indic font collection)
FONT_NAME = "Helvetica"
FONT_BOLD_NAME = "Helvetica-Bold"

nirmala_path = Config.NIRMALA_FONT_PATH
if nirmala_path and os.path.exists(nirmala_path):
    try:
        pdfmetrics.registerFont(TTFont('Nirmala', nirmala_path, subfontIndex=0))
        pdfmetrics.registerFont(TTFont('Nirmala-Bold', nirmala_path, subfontIndex=1))
        FONT_NAME = "Nirmala"
        FONT_BOLD_NAME = "Nirmala-Bold"
        logger.info("[DocumentGenerator] Successfully registered Nirmala and Nirmala-Bold fonts from TTC.")
    except Exception as e:
        logger.warning(f"[DocumentGenerator] Failed to register Nirmala font collection: {e}")

class DocumentGenerator:
    def _get_safe_filename(self, title: str) -> str:
        safe_title = "".join(c for c in title if c.isalnum() or c in (" ", "-", "_")).strip()
        return safe_title or "Video_Summary"

    def generate_markdown(self, summary_data: Dict[str, Any], meta: Dict[str, Any], output_path: str) -> str:
        title = meta.get("title", "Video Summary")
        
        md_lines = [
            f"# AI Video Summary: {title}",
            "",
            f"**Source**: {meta.get('source_url', 'N/A')} ({meta.get('source_type', 'Upload')})",
            f"**Duration**: {meta.get('duration_str', 'N/A')}",
            f"**Generated Time**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "## Executive Summary",
            summary_data.get("executive_summary", "N/A"),
            "",
            "## Detailed Summary",
            summary_data.get("detailed_summary", "N/A"),
            "",
            "## Key Topics",
            ""
        ]
        
        for topic in summary_data.get("key_topics", []):
            md_lines.append(f"- **{topic.get('topic')}**: {topic.get('description')}")
            
        md_lines.extend([
            "",
            "## Important Highlights",
            ""
        ])
        
        for point in summary_data.get("important_points", []):
            md_lines.append(f"- {point}")
            
        md_lines.extend([
            "",
            "## Action Items",
            ""
        ])
        
        for item in summary_data.get("action_items", []):
            md_lines.append(f"- {item}")
            
        md_lines.extend([
            "",
            "## Timeline",
            ""
        ])
        
        for t in summary_data.get("timeline", []):
            md_lines.append(f"- `[{t.get('timestamp')}]` {t.get('event')}")
            
        md_lines.extend([
            "",
            "## Conclusion",
            summary_data.get("conclusion", "N/A")
        ])
        
        content = "\n".join(md_lines)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
            
        logger.info(f"[DocumentGenerator] Markdown summary generated: {output_path}")
        return output_path

    def generate_docx(self, summary_data: Dict[str, Any], meta: Dict[str, Any], output_path: str) -> str:
        title = meta.get("title", "Video Summary")
        
        doc = DocxDocument()
        
        # Style Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"AI Video Summary: {title}")
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = RGBColor(26, 82, 118) # Navy Blue
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Metadata
        doc.add_paragraph(f"Source: {meta.get('source_url', 'N/A')} ({meta.get('source_type', 'Upload')})")
        doc.add_paragraph(f"Duration: {meta.get('duration_str', 'N/A')}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("").paragraph_format.space_after = Pt(12)
        
        # Executive Summary
        h_exec = doc.add_heading("Executive Summary", level=2)
        h_exec.runs[0].font.color.rgb = RGBColor(26, 82, 118)
        doc.add_paragraph(summary_data.get("executive_summary", ""))
        
        # Detailed Summary
        h_det = doc.add_heading("Detailed Summary", level=2)
        h_det.runs[0].font.color.rgb = RGBColor(26, 82, 118)
        doc.add_paragraph(summary_data.get("detailed_summary", ""))
        
        # Key Topics
        h_top = doc.add_heading("Key Topics", level=2)
        h_top.runs[0].font.color.rgb = RGBColor(26, 82, 118)
        for topic in summary_data.get("key_topics", []):
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(f"{topic.get('topic')}: ")
            r.bold = True
            p.add_run(topic.get('description'))
            
        # Important Points
        h_points = doc.add_heading("Important Highlights", level=2)
        h_points.runs[0].font.color.rgb = RGBColor(26, 82, 118)
        for point in summary_data.get("important_points", []):
            doc.add_paragraph(point, style='List Bullet')
            
        # Action Items
        h_act = doc.add_heading("Action Items", level=2)
        h_act.runs[0].font.color.rgb = RGBColor(26, 82, 118)
        for item in summary_data.get("action_items", []):
            doc.add_paragraph(item, style='List Bullet')
            
        # Timeline
        h_time = doc.add_heading("Timeline", level=2)
        h_time.runs[0].font.color.rgb = RGBColor(26, 82, 118)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Timestamp'
        hdr_cells[1].text = 'Event / Topic Shift'
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True
        
        for t in summary_data.get("timeline", []):
            row_cells = table.add_row().cells
            row_cells[0].text = f"[{t.get('timestamp')}]"
            row_cells[1].text = t.get('event')
            
        doc.add_paragraph("")
        
        # Conclusion
        h_con = doc.add_heading("Conclusion", level=2)
        h_con.runs[0].font.color.rgb = RGBColor(26, 82, 118)
        doc.add_paragraph(summary_data.get("conclusion", ""))
        
        doc.save(output_path)
        logger.info(f"[DocumentGenerator] Word summary generated: {output_path}")
        return output_path

    def generate_pdf(self, summary_data: Dict[str, Any], meta: Dict[str, Any], output_path: str) -> str:
        title = meta.get("title", "Video Summary")
        
        # Initialize SimpleDocTemplate
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        
        # Define Premium Color Palette
        navy = HexColor('#1A5276')
        charcoal = HexColor('#2C3E50')
        slate = HexColor('#7F8C8D')
        light_bg = HexColor('#F8F9F9')
        
        # Custom styles
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Normal'],
            fontName=FONT_BOLD_NAME,
            fontSize=22,
            leading=26,
            textColor=navy,
            spaceAfter=12
        )
        
        meta_label_style = ParagraphStyle(
            'MetaLabel',
            parent=styles['Normal'],
            fontName=FONT_BOLD_NAME,
            fontSize=10,
            leading=14,
            textColor=charcoal
        )
        
        meta_val_style = ParagraphStyle(
            'MetaVal',
            parent=styles['Normal'],
            fontName=FONT_NAME,
            fontSize=10,
            leading=14,
            textColor=slate
        )
        
        h2_style = ParagraphStyle(
            'Heading2',
            parent=styles['Heading2'],
            fontName=FONT_BOLD_NAME,
            fontSize=14,
            leading=18,
            textColor=navy,
            spaceBefore=16,
            spaceAfter=8,
            keepWithNext=True
        )
        
        body_style = ParagraphStyle(
            'Body',
            parent=styles['BodyText'],
            fontName=FONT_NAME,
            fontSize=10,
            leading=14,
            textColor=charcoal,
            spaceAfter=10
        )
        
        bullet_style = ParagraphStyle(
            'Bullet',
            parent=body_style,
            leftIndent=20,
            firstLineIndent=-10,
            spaceAfter=6
        )
        
        table_cell_style = ParagraphStyle(
            'TableCell',
            parent=styles['Normal'],
            fontName=FONT_NAME,
            fontSize=9.5,
            leading=13,
            textColor=charcoal
        )
        
        table_hdr_style = ParagraphStyle(
            'TableHdr',
            parent=styles['Normal'],
            fontName=FONT_BOLD_NAME,
            fontSize=10,
            leading=14,
            textColor=colors.white
        )

        story = []
        
        # Title
        story.append(Paragraph(f"AI Video Summary: {title}", title_style))
        story.append(Spacer(1, 8))
        
        # Metadata Table
        meta_data = [
            [Paragraph("Source URL:", meta_label_style), Paragraph(meta.get('source_url', 'N/A'), meta_val_style)],
            [Paragraph("Provider:", meta_label_style), Paragraph(meta.get('source_type', 'Upload'), meta_val_style)],
            [Paragraph("Video Duration:", meta_label_style), Paragraph(meta.get('duration_str', 'N/A'), meta_val_style)],
            [Paragraph("Report Generated:", meta_label_style), Paragraph(datetime.now().strftime('%Y-%m-%d %H:%M:%S'), meta_val_style)]
        ]
        meta_table = Table(meta_data, colWidths=[110, 390])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), light_bg),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('LEFTPADDING', (0,0), (-1,-1), 8),
            ('RIGHTPADDING', (0,0), (-1,-1), 8),
            ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.lightgrey)
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 15))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", h2_style))
        story.append(Paragraph(summary_data.get("executive_summary", "N/A"), body_style))
        
        # Detailed Summary
        story.append(Paragraph("Detailed Summary", h2_style))
        story.append(Paragraph(summary_data.get("detailed_summary", "N/A"), body_style))
        
        # Key Topics
        story.append(Paragraph("Key Topics", h2_style))
        for topic in summary_data.get("key_topics", []):
            topic_p = f"<b>{topic.get('topic')}</b>: {topic.get('description')}"
            story.append(Paragraph(f"&bull; {topic_p}", bullet_style))
            
        # Important Highlights
        story.append(Paragraph("Important Highlights", h2_style))
        for point in summary_data.get("important_points", []):
            story.append(Paragraph(f"&bull; {point}", bullet_style))
            
        # Action Items
        story.append(Paragraph("Action Items", h2_style))
        for item in summary_data.get("action_items", []):
            story.append(Paragraph(f"&bull; {item}", bullet_style))
            
        # Timeline
        story.append(Paragraph("Timeline", h2_style))
        timeline_data = [[
            Paragraph("Timestamp", table_hdr_style),
            Paragraph("Topic shift / Discussion point", table_hdr_style)
        ]]
        for t in summary_data.get("timeline", []):
            timeline_data.append([
                Paragraph(f"<b>[{t.get('timestamp')}]</b>", table_cell_style),
                Paragraph(t.get('event'), table_cell_style)
            ])
            
        timeline_table = Table(timeline_data, colWidths=[90, 414])
        timeline_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), navy),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, light_bg]),
            ('GRID', (0,0), (-1,-1), 0.5, colors.lightgrey),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('LEFTPADDING', (0,0), (-1,-1), 8),
            ('RIGHTPADDING', (0,0), (-1,-1), 8),
        ]))
        story.append(timeline_table)
        story.append(Spacer(1, 15))
        
        # Conclusion
        story.append(Paragraph("Conclusion", h2_style))
        story.append(Paragraph(summary_data.get("conclusion", "N/A"), body_style))
        
        # Build Document
        doc.build(story)
        logger.info(f"[DocumentGenerator] PDF summary generated: {output_path}")
        return output_path

    def generate_all(self, summary_data: Dict[str, Any], meta: Dict[str, Any], user_id: int, base_dir: str) -> Dict[str, str]:
        """Generates PDF, DOCX, and MD summaries and returns dict of their final filepaths."""
        user_folder = os.path.join(base_dir, str(user_id))
        os.makedirs(user_folder, exist_ok=True)
        
        safe_name = self._get_safe_filename(meta.get("title", "Video_Summary"))
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        md_name = f"{safe_name}_{timestamp}.md"
        docx_name = f"{safe_name}_{timestamp}.docx"
        pdf_name = f"{safe_name}_{timestamp}.pdf"
        
        md_path = os.path.join(user_folder, md_name)
        docx_path = os.path.join(user_folder, docx_name)
        pdf_path = os.path.join(user_folder, pdf_name)
        
        self.generate_markdown(summary_data, meta, md_path)
        self.generate_docx(summary_data, meta, docx_path)
        self.generate_pdf(summary_data, meta, pdf_path)
        
        return {
            "md": {"name": md_name, "path": md_path},
            "docx": {"name": docx_name, "path": docx_path},
            "pdf": {"name": pdf_name, "path": pdf_path}
        }